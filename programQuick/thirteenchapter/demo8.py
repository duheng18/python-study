import docx

doc = docx.Document('demo.docx')

# Document Title
# print(doc.paragraphs[0].text)

# _ParagraphStyle('Title') id: 4400855080
# print(doc.paragraphs[0].style)

doc.paragraphs[0].style = 'Normal'

# A plain paragraph with some bold and some italic
# print(doc.paragraphs[1].text)

# A plain paragraph with
print(doc.paragraphs[1].runs[0].text)

#  some
print(doc.paragraphs[1].runs[1].text)

# bold
print(doc.paragraphs[1].runs[2].text)

#  and some
print(doc.paragraphs[1].runs[3].text)

doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('restyled.docx')
