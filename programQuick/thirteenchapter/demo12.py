import docx
from docx.shared import Inches, Cm

doc = docx.Document()
doc.add_paragraph('This is on the first page!')
# 要添加换行符(而不是开始一个新的段落)，可以在 Run 对象上调用 add_break()方法，
# 换行符将出现在它后面。如果希望添加换页符，可以将 docx.text.WD_BREAK.PAGE
# 作为唯一的参数，传递给 add_break()，就像下面代码中间所做的一样.
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.add_picture('zophie.png', width=Inches(1), height=Cm(4))
doc.save('twoPage.docx')
