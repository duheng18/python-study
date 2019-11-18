import docx

doc = docx.Document('demo.docx')
# 7
# print(len(doc.paragraphs))

# Document Title
# print(doc.paragraphs[0].text)

# A plain paragraph with some bold and some italic
# print(doc.paragraphs[1].text)

# 5
# print(len(doc.paragraphs[1].runs))

# A plain paragraph with
# print(doc.paragraphs[1].runs[0].text)

#  some
# print(doc.paragraphs[1].runs[1].text)

# bold
# print(doc.paragraphs[1].runs[2].text)

#  and some
# print(doc.paragraphs[1].runs[3].text)

# italic
print(doc.paragraphs[1].runs[4].text)
